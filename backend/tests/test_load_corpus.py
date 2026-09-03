"""The promises of scripts/dev/build_load_corpus.py, held against small runs.

The generator exists for one run on rented hardware that takes the better part
of a day, and everything about that run has to be repeatable afterwards: the
same seed has to produce the same 50.000 files, or the measurement report is a
story rather than a record. This suite proves that on 50 and on 500 files and
never on the full set. A test that writes 20 GB is not a test, it is the load
run with an assertion bolted on.

The module under test is a script and not part of the installed package, so it
is loaded from its path. That is deliberate: it belongs next to build_corpus.py
where a reader looks for corpus tooling, and it must stay runnable with a bare
``uv run python scripts/dev/build_load_corpus.py``.
"""

from __future__ import annotations

import importlib.util
import io
import sys
import tracemalloc
from collections.abc import Iterator
from pathlib import Path
from types import ModuleType
from typing import Any

import pytest

REPO_ROOT = Path(__file__).resolve().parents[2]
GENERATOR_PATH = REPO_ROOT / "scripts" / "dev" / "build_load_corpus.py"


def _load_generator() -> ModuleType:
    specification = importlib.util.spec_from_file_location("build_load_corpus", GENERATOR_PATH)
    assert specification is not None
    assert specification.loader is not None
    module = importlib.util.module_from_spec(specification)
    sys.modules[specification.name] = module
    specification.loader.exec_module(module)
    return module


corpus = _load_generator()

# The two sizes this suite works with. 50 is the fast one that every structural
# assertion uses, 500 is the size of the dry run that goes on the box before the
# full run, and it is the size the reproducibility claim is made at.
SMALL = 50
DRY_RUN = 500


def _run(out: Path, seed: str, files: int) -> tuple[Any, str]:
    """One generator run into ``out``, with the report captured as a string."""
    report = io.StringIO()
    summary = corpus.generate(out=out, seed=seed, files=files, report=report)
    return summary, report.getvalue()


@pytest.fixture(scope="module")
def dry_run_pair(tmp_path_factory: pytest.TempPathFactory) -> tuple[Any, Any, Any]:
    """Two runs with the same seed and one with another, all at 500 files."""
    first, _ = _run(tmp_path_factory.mktemp("first"), "phase5-dry", DRY_RUN)
    second, _ = _run(tmp_path_factory.mktemp("second"), "phase5-dry", DRY_RUN)
    other, _ = _run(tmp_path_factory.mktemp("other"), "phase5-other", DRY_RUN)
    return first, second, other


def test_the_same_seed_writes_the_same_bytes(tmp_path: Path) -> None:
    """Bitwise, not "the same kind of file": every byte of every file."""
    left = tmp_path / "left"
    right = tmp_path / "right"
    first, _ = _run(left, "phase5", SMALL)
    second, _ = _run(right, "phase5", SMALL)

    left_files = sorted(path.name for path in left.iterdir())
    right_files = sorted(path.name for path in right.iterdir())
    assert left_files == right_files
    assert len(left_files) == SMALL
    for name in left_files:
        assert (left / name).read_bytes() == (right / name).read_bytes(), name
    assert first.checksum == second.checksum


def test_the_same_seed_gives_the_same_checksum_at_the_dry_run_size(
    dry_run_pair: tuple[Any, Any, Any],
) -> None:
    first, second, _ = dry_run_pair
    assert first.checksum == second.checksum
    assert first.total_bytes == second.total_bytes
    assert first.files == DRY_RUN


def test_another_seed_gives_another_checksum(dry_run_pair: tuple[Any, Any, Any]) -> None:
    first, _, other = dry_run_pair
    assert first.checksum != other.checksum


def test_the_dry_run_has_the_shares_of_the_full_run() -> None:
    """The 500 files are the 50.000 in miniature, and no category falls out."""
    full = corpus.allocate(50_000)
    dry = corpus.allocate(DRY_RUN)

    assert sum(full.values()) == 50_000
    assert sum(dry.values()) == DRY_RUN
    assert set(full) == set(dry)
    for key, count in dry.items():
        assert count >= 1, key
        # One file out of 500 is 0,2 percent, so a category may miss its share
        # by up to that much and still be the same distribution.
        assert abs(count / DRY_RUN - full[key] / 50_000) <= 0.003, key


def test_exactly_twenty_files_lie_above_the_size_cap() -> None:
    """The oversize category is the only one above the cap, and it holds 20."""
    from findling.config import MAX_FILE_BYTES

    allocation = corpus.allocate(50_000)
    above = {
        category.key: allocation[category.key]
        for category in corpus.CATEGORIES
        if category.target_bytes > MAX_FILE_BYTES
    }
    assert above == {corpus.OVERSIZE_KEY: 20}


def test_the_dry_run_report_counts_what_the_allocation_promises(tmp_path: Path) -> None:
    from findling.config import MAX_FILE_BYTES

    summary, report = _run(tmp_path / "corpus", "phase5-report", SMALL)
    rows = [line.split(",") for line in report.strip().splitlines()]
    assert rows[0] == ["name", "bytes", "sha256"]
    body = rows[1:]
    assert len(body) == SMALL
    assert summary.total_bytes == sum(int(row[1]) for row in body)

    oversize = [row for row in body if int(row[1]) > MAX_FILE_BYTES]
    assert len(oversize) == corpus.allocate(SMALL)[corpus.OVERSIZE_KEY]
    written = sorted(path.name for path in (tmp_path / "corpus").iterdir())
    assert written == sorted(row[0] for row in body)


def test_memory_does_not_follow_the_size_of_the_written_file(tmp_path: Path) -> None:
    """The generator streams, so one very large file costs no more than a chunk.

    This is the property that makes 20 GB possible at all: build_corpus.py holds
    every payload in a dict, and that approach cannot be scaled by a factor of
    six hundred.
    """
    from findling.config import MAX_FILE_BYTES

    target = tmp_path / "huge.csv"
    rng = corpus.Rng("phase5", "oversize", 1)
    tracemalloc.start()
    try:
        size, _ = corpus.write_file(target, corpus.build_oversize(rng, "csv", corpus.OVERSIZE_BYTES))
        _, peak = tracemalloc.get_traced_memory()
    finally:
        tracemalloc.stop()

    assert size > MAX_FILE_BYTES
    assert peak < 8 * 1024 * 1024, f"peak {peak} for a file of {size} bytes"


def test_a_moved_font_stops_the_run_before_the_first_byte(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> None:
    out = tmp_path / "corpus"
    monkeypatch.setattr(corpus, "DEJAVU_SANS_SHA256", "0" * 64)
    corpus.load_font.cache_clear()
    try:
        with pytest.raises(SystemExit):
            _run(out, "phase5", SMALL)
    finally:
        corpus.load_font.cache_clear()
    assert not out.exists() or list(out.iterdir()) == []


def test_an_abort_leaves_nothing_that_looks_finished(tmp_path: Path) -> None:
    target = tmp_path / "07000-bescheid.pdf"

    def chunks() -> Iterator[bytes]:
        yield b"the first half of a file"
        message = "the run was killed here"
        raise RuntimeError(message)

    with pytest.raises(RuntimeError):
        corpus.write_file(target, chunks())

    assert not target.exists()
    assert list(tmp_path.iterdir()) == []


def test_the_command_prints_seed_files_bytes_and_checksum(
    tmp_path: Path, capsys: pytest.CaptureFixture[str]
) -> None:
    out = tmp_path / "corpus"
    report = tmp_path / "report.csv"
    code = corpus.main(
        ["--seed", "phase5-cli", "--files", str(SMALL), "--out", str(out), "--report", str(report)]
    )
    printed = capsys.readouterr().out

    assert code == 0
    assert "seed=phase5-cli" in printed
    assert f"files={SMALL}" in printed
    assert "bytes=" in printed
    assert "checksum=" in printed
    assert report.read_text(encoding="utf-8").startswith("name,bytes,sha256")


def test_the_dry_run_shortcut_selects_five_hundred_files() -> None:
    arguments = corpus.parse_arguments(["--seed", "x", "--out", "somewhere", "--dry-run-files"])
    assert arguments.files == DRY_RUN


def test_the_generated_documents_are_readable_by_the_extractors(tmp_path: Path) -> None:
    """A corpus of unreadable files would measure the error path for a day.

    Every format the generator writes is opened here with the same library the
    backend uses for it, so a broken part shows up in a second rather than after
    twenty hours on rented hardware.
    """
    import docx
    import openpyxl
    import pptx
    import pypdfium2

    from findling.extract.odf import extract_odf

    out = tmp_path / "corpus"
    _run(out, "phase5-read", SMALL)
    seen: set[str] = set()
    for path in sorted(out.iterdir()):
        suffix = path.suffix.lstrip(".")
        seen.add(suffix)
        if suffix == "pdf":
            document = pypdfium2.PdfDocument(str(path))
            assert len(document) >= 1
            document.close()
        elif suffix == "docx":
            assert any(paragraph.text for paragraph in docx.Document(str(path)).paragraphs)
        elif suffix == "xlsx":
            workbook = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            try:
                rows = next(workbook.worksheets[0].iter_rows(values_only=True))
            finally:
                workbook.close()
            assert any(cell for cell in rows)
        elif suffix == "pptx":
            slides = pptx.Presentation(str(path)).slides
            assert len(slides) >= 1
        elif suffix in {"odt", "ods"}:
            outcome = extract_odf(str(path))
            assert outcome.text, path.name

    assert {"pdf", "txt"} <= seen


def test_neither_new_file_carries_a_dash_or_an_indeterminate_source() -> None:
    """Two house rules, kept by a test instead of by attention.

    The dashes are the typography rule of this project. The three names are the
    reproducibility rule: a corpus that draws from the clock or from the kernel
    entropy pool cannot be rebuilt, and the seed on the report would be a
    decoration. The names are assembled rather than written out, because this
    file would otherwise fail its own second half.
    """
    for path in (GENERATOR_PATH, Path(__file__)):
        text = path.read_text(encoding="utf-8")
        assert "—" not in text, path.name
        assert "–" not in text, path.name

    generator = GENERATOR_PATH.read_text(encoding="utf-8")
    for name in ("ur" + "andom", "random." + "random", "time." + "time"):
        assert name not in generator, name
